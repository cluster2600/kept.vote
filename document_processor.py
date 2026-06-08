"""Document text extraction for PDF and Word files.

:class:`DocumentProcessor` accepts raw file bytes and returns extracted plain
text. Extraction libraries (PyPDF2, python-docx) are synchronous and CPU-bound,
so the async entrypoint offloads them to a worker thread to avoid blocking the
event loop.
"""

from __future__ import annotations

import io

from docx import Document as DocxDocument
from PyPDF2 import PdfReader
from starlette.concurrency import run_in_threadpool

# MIME types and extensions we know how to handle.
PDF_CONTENT_TYPES = {"application/pdf"}
DOCX_CONTENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}


class UnsupportedDocumentError(ValueError):
    """Raised when a file's type is not a supported PDF/Word document."""


class DocumentExtractionError(RuntimeError):
    """Raised when a supported file cannot be parsed (corrupt/encrypted)."""


class DocumentProcessor:
    """Extracts plain text from uploaded PDF and Word documents."""

    @staticmethod
    def _extract_pdf(data: bytes) -> str:
        """Extract and concatenate text from every page of a PDF.

        Raises :class:`DocumentExtractionError` if the file cannot be read.
        """
        try:
            reader = PdfReader(io.BytesIO(data))
        except Exception as exc:  # noqa: BLE001 - normalize to our error type
            raise DocumentExtractionError(f"Could not read PDF: {exc}") from exc

        parts: list[str] = []
        for page in reader.pages:
            # extract_text() can return None for image-only pages.
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text)
        return "\n\n".join(parts).strip()

    @staticmethod
    def _extract_docx(data: bytes) -> str:
        """Extract text from a Word document's paragraphs and tables."""
        try:
            document = DocxDocument(io.BytesIO(data))
        except Exception as exc:  # noqa: BLE001 - normalize to our error type
            raise DocumentExtractionError(
                f"Could not read Word document: {exc}"
            ) from exc

        parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]

        # Pull text out of tables too — manifestos and reports lean on them.
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts).strip()

    @classmethod
    def _extract_sync(cls, data: bytes, content_type: str, filename: str) -> str:
        """Dispatch to the right extractor based on content type/extension."""
        ctype = (content_type or "").lower()
        name = (filename or "").lower()

        if ctype in PDF_CONTENT_TYPES or name.endswith(".pdf"):
            return cls._extract_pdf(data)
        if ctype in DOCX_CONTENT_TYPES or name.endswith((".docx", ".doc")):
            return cls._extract_docx(data)

        raise UnsupportedDocumentError(
            "Unsupported document type "
            f"(content_type={content_type!r}, filename={filename!r}). "
            "Only PDF and Word (.docx) files are supported."
        )

    async def extract_text(
        self,
        data: bytes,
        content_type: str,
        filename: str,
    ) -> str:
        """Extract text from ``data`` without blocking the event loop.

        Args:
            data: Raw file bytes.
            content_type: The upload's MIME type (e.g. ``application/pdf``).
            filename: Original filename, used as an extension fallback.

        Returns:
            The extracted plain text (may be empty for image-only documents).

        Raises:
            UnsupportedDocumentError: If the file is not a PDF/Word document.
            DocumentExtractionError: If a supported file cannot be parsed.
        """
        return await run_in_threadpool(
            self._extract_sync, data, content_type, filename
        )
